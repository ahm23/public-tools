#!/usr/bin/env python3
"""
Requirements specification converter — DOCX <-> JSON <-> XLSX.

Bridges the Requirements Explorer app (Electrobun) to Word/Excel files.
The app spawns this script and talks to it over files + a one-line JSON
result on stdout.

Internal JSON format (the app's native format):
    {
      "format": "requirements-explorer-spec",
      "version": 1,
      "blocks": [
        {"type": "heading", "level": 1, "title": "..."},
        {"type": "item", "kind": "requirement" | "comment", "id": "REQ4510", "text": "..."}
      ]
    }

Rules (mirroring the legacy req-docx_to_xlsx.py):
  * Title is only meaningful on heading rows — items have a blank title.
  * An item is a "requirement" when it carries an external ID token
    ([AAAA####]) or its text contains shall/should/may; otherwise "comment".
  * Heading nesting levels survive in DOCX and JSON; XLSX flattens them
    (the XLSX format has no level column).

Subcommands:
    req_convert.py docx2json  input.docx  [output.json]
    req_convert.py json2docx  input.json   output.docx
    req_convert.py xlsx2json  input.xlsx  [output.json]
    req_convert.py json2xlsx  input.json   output.xlsx

Usage from the app:
    <venv>/bin/python req_convert.py docx2json in.docx out.json
    # stdout: {"ok": true, "headings": 3, "requirements": 8, "comments": 2, "output": "out.json"}
"""

import argparse
import json
import re
import sys
from pathlib import Path

# --------------------------------------------------------------------------- #
# DOCX parsing helpers (shared with the legacy converter)
# --------------------------------------------------------------------------- #

TOKEN_RE = re.compile(r"\[([A-Z]{2,6}\d{2,6})\]")
SHALL_WORDS = ("shall", "should", "may")


def _qn(name: str) -> str:
    return "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + name


def get_num_info(paragraph):
    """Return (numId, ilvl) for a paragraph, or (None, None) if not numbered."""
    pPr = paragraph._element.find(_qn("pPr"))
    if pPr is None:
        return None, None
    numPr = pPr.find(_qn("numPr"))
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(_qn("ilvl"))
    numId_el = numPr.find(_qn("numId"))
    ilvl = int(ilvl_el.get(_qn("val"))) if ilvl_el is not None else None
    numId = int(numId_el.get(_qn("val"))) if numId_el is not None else None
    return numId, ilvl


def has_drawing(paragraph) -> bool:
    return len(paragraph._element.findall(f".//{_qn('drawing')}")) > 0


def get_bold_text(paragraph) -> str:
    style_bold = False
    if paragraph.style and paragraph.style.font:
        style_bold = paragraph.style.font.bold is True
    parts = []
    for r in paragraph.runs:
        if (r.bold is True) or (r.font.bold is True) or style_bold:
            parts.append(r.text)
    return "".join(parts).strip()


def is_figure_or_table_caption(text: str, bold_text: str) -> bool:
    if not bold_text:
        return False
    t = text.strip()
    return t.startswith("Figure") or t.startswith("Table")


def get_clean_text(paragraph) -> str:
    text = "".join(r.text for r in paragraph.runs).strip()
    if has_drawing(paragraph):
        text = f"{text} [IMAGE]" if text else "[IMAGE]"
    return text


def is_heading(paragraph):
    """(True, level) for Word built-in Heading 1..9 styles, else (False, None)."""
    if paragraph.style and paragraph.style.name:
        name = paragraph.style.name
        if name.startswith("Heading"):
            try:
                level = int("".join(filter(str.isdigit, name)))
                if 1 <= level <= 9:
                    return True, level
            except ValueError:
                pass
    return False, None


def classify_item(text: str, ext_id: str) -> str:
    """Return 'requirement' or 'comment'."""
    if ext_id:
        return "requirement"
    lowered = text.lower()
    return "requirement" if any(w in lowered for w in SHALL_WORDS) else "comment"


# --------------------------------------------------------------------------- #
# DOCX -> JSON
# --------------------------------------------------------------------------- #

def split_text_to_items(text: str):
    """Split body text on [AAAA####] tokens into item blocks."""
    items = []
    matches = list(TOKEN_RE.finditer(text))
    if not matches:
        items.append({"type": "item", "kind": classify_item(text, ""),
                      "id": "", "text": text.strip()})
        return items

    leading = text[: matches[0].start()].strip()
    if leading:
        items.append({"type": "item", "kind": classify_item(leading, ""),
                      "id": "", "text": leading})

    for i, m in enumerate(matches):
        ext_id = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else None
        body = text[start:end].strip() if end is not None else text[start:].strip()
        items.append({"type": "item", "kind": "requirement",
                      "id": ext_id, "text": body})
    return items


def parse_docx_to_blocks(docx_path: Path):
    from docx import Document
    from docx.text.paragraph import Paragraph

    doc = Document(str(docx_path))
    blocks = []
    body_parts = []          # raw body lines accumulated until next heading
    seen_heading = False

    def flush_body():
        nonlocal body_parts
        text = "\n".join(body_parts).strip()
        if text:
            blocks.extend(split_text_to_items(text))
        body_parts = []

    for child in doc.element.body.iterchildren():
        if child.tag == _qn("p"):
            para = Paragraph(child, doc)
            num_id, ilvl = get_num_info(para)
            clean_text = get_clean_text(para)
            bold_text = get_bold_text(para)

            if not clean_text and not has_drawing(para) and num_id is None:
                continue
            if is_figure_or_table_caption(clean_text, bold_text):
                continue

            is_hd, hd_level = is_heading(para)
            if is_hd:
                flush_body()
                blocks.append({"type": "heading", "level": hd_level,
                               "title": clean_text})
                seen_heading = True
                continue

            if num_id is not None:
                body_parts.append(f"- {clean_text}")
                continue
            if clean_text:
                body_parts.append(clean_text)
                continue
            body_parts.append(clean_text)

        elif child.tag == _qn("tbl"):
            if not body_parts and not seen_heading:
                # Table before any heading becomes its own heading (legacy behavior)
                blocks.append({"type": "heading", "level": 1, "title": "[TABLE]"})
                seen_heading = True
            body_parts.append("[TABLE]")

    flush_body()
    return blocks


def docx2json(input_path: Path, output_path: Path) -> dict:
    blocks = parse_docx_to_blocks(input_path)
    spec = {"format": "requirements-explorer-spec", "version": 1, "blocks": blocks}
    _write_json(spec, output_path)
    return _counts(spec)


# --------------------------------------------------------------------------- #
# JSON -> DOCX
# --------------------------------------------------------------------------- #

def json2docx(input_path: Path, output_path: Path) -> dict:
    from docx import Document

    spec = _read_json(input_path)
    doc = Document()

    for b in spec.get("blocks", []):
        if b.get("type") == "heading":
            level = min(max(int(b.get("level", 1)), 1), 9)
            doc.add_heading(b.get("title", ""), level=level)
        else:
            text = (b.get("text") or "").strip()
            ext_id = (b.get("id") or "").strip()
            if ext_id:
                doc.add_paragraph(f"[{ext_id}] {text}")
            elif text:
                doc.add_paragraph(text)

    doc.save(str(output_path))
    return _counts(spec)


# --------------------------------------------------------------------------- #
# XLSX <-> JSON
# --------------------------------------------------------------------------- #

XLSX_HEADERS = ["Title", "Type", "External ID", "Text"]


def xlsx2json(input_path: Path, output_path: Path) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(str(input_path), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))

    header_idx = None
    for i, row in enumerate(rows):
        if row and str(row[0]).strip() == "Title" and len(row) >= 4:
            header_idx = i
            break
    if header_idx is None:
        raise ValueError("XLSX does not look like a converted spec (missing header row)")

    blocks = []
    for row in rows[header_idx + 1:]:
        if not row or all(c is None or str(c).strip() == "" for c in row):
            continue
        title = str(row[0]).strip() if row[0] is not None else ""
        rtype = str(row[1]).strip() if row[1] is not None else ""
        ext_id = str(row[2]).strip() if row[2] is not None else ""
        text = str(row[3]).strip() if row[3] is not None else ""

        if rtype == "Heading" or (title and not ext_id and not text):
            blocks.append({"type": "heading", "level": 1, "title": title or text})
        else:
            kind = "requirement" if (ext_id or rtype == "Requirement"
                                     or classify_item(text, "") == "requirement") else "comment"
            blocks.append({"type": "item", "kind": kind, "id": ext_id, "text": text})

    wb.close()
    spec = {"format": "requirements-explorer-spec", "version": 1, "blocks": blocks}
    _write_json(spec, output_path)
    return _counts(spec)


def json2xlsx(input_path: Path, output_path: Path) -> dict:
    from openpyxl import Workbook

    spec = _read_json(input_path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Specification"
    ws.append(XLSX_HEADERS)

    for b in spec.get("blocks", []):
        if b.get("type") == "heading":
            title = b.get("title", "")
            ws.append([title, "Heading", "", title])
        else:
            text = (b.get("text") or "").strip()
            ext_id = (b.get("id") or "").strip()
            rtype = "Requirement" if (ext_id or classify_item(text, "") == "requirement") else "Comment"
            ws.append(["", rtype, ext_id, text])

    ws.column_dimensions["A"].width = 40
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 80

    wb.save(str(output_path))
    return _counts(spec)


# --------------------------------------------------------------------------- #
# Shared
# --------------------------------------------------------------------------- #

def _read_json(path: Path) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        spec = json.load(f)
    if not isinstance(spec, dict) or "blocks" not in spec:
        raise ValueError(f"{path} is not a requirements-explorer spec file")
    return spec


def _write_json(spec: dict, path: Path) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(spec, f, ensure_ascii=False, indent=2)


def _counts(spec: dict) -> dict:
    headings = sum(1 for b in spec.get("blocks", []) if b.get("type") == "heading")
    requirements = sum(1 for b in spec.get("blocks", [])
                       if b.get("type") == "item" and b.get("kind") == "requirement")
    comments = sum(1 for b in spec.get("blocks", [])
                   if b.get("type") == "item" and b.get("kind") != "requirement")
    return {"ok": True, "headings": headings, "requirements": requirements,
            "comments": comments, "blocks": headings + requirements + comments}


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        prog="req_convert.py",
        description="Convert requirement specifications between DOCX, JSON, and XLSX.")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("docx2json", help="DOCX -> internal JSON")
    p.add_argument("input", help="input .docx")
    p.add_argument("output", nargs="?", help="output .json (default: input + .json)")

    p = sub.add_parser("json2docx", help="internal JSON -> DOCX")
    p.add_argument("input", help="input .json")
    p.add_argument("output", help="output .docx")

    p = sub.add_parser("xlsx2json", help="converted spec XLSX -> internal JSON")
    p.add_argument("input", help="input .xlsx")
    p.add_argument("output", nargs="?", help="output .json (default: input + .json)")

    p = sub.add_parser("json2xlsx", help="internal JSON -> spec XLSX (4 columns)")
    p.add_argument("input", help="input .json")
    p.add_argument("output", help="output .xlsx")

    args = ap.parse_args(argv)
    inp = Path(args.input).expanduser()
    if not inp.exists():
        print(json.dumps({"ok": False, "error": f"input file not found: {inp}"}))
        return 1

    try:
        if args.cmd == "docx2json":
            out = Path(args.output).expanduser() if args.output else inp.with_suffix(".json")
            result = docx2json(inp, out)
        elif args.cmd == "json2docx":
            out = Path(args.output).expanduser()
            result = json2docx(inp, out)
        elif args.cmd == "xlsx2json":
            out = Path(args.output).expanduser() if args.output else inp.with_suffix(".json")
            result = xlsx2json(inp, out)
        else:
            out = Path(args.output).expanduser()
            result = json2xlsx(inp, out)
        result["output"] = str(out)
    except Exception as e:  # noqa: BLE001 — report any failure as JSON
        print(json.dumps({"ok": False, "error": f"{type(e).__name__}: {e}"}))
        return 1

    print(json.dumps(result))
    return 0


if __name__ == "__main__":
    sys.exit(main())
